import os
import shutil
import json
from datetime import datetime
import fitz  # PyMuPDF
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from database import SessionLocal, Document, CrossAnalysis, AnalysisPreset
from extractor import process_pdf_with_vision, perform_cross_analysis
from sqlalchemy.orm import Session
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
import io
import re
import pandas as pd
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
import zipfile
import tempfile
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.shared import qn, OxmlElement
import docx.opc.constants
import docx

app = FastAPI(title="Cloud File Extractor API")

# Configure CORS for local UI access
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.middleware("http")
async def log_requests(request: Request, call_next):
    print(f"[RETE] Richiesta in arrivo: {request.method} {request.url.path}")
    response = await call_next(request)
    return response

def create_docx_from_markdown(md_text: str):
    """Simple parser for legal markdown to professional DOCX"""
    doc = DocxDocument()
    
    # Custom Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)

    lines = md_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        # Headers
        if line.startswith('### '):
            h = doc.add_paragraph(line.replace('### ', ''), style='Heading 3')
        elif line.startswith('## '):
            h = doc.add_paragraph(line.replace('## ', ''), style='Heading 2')
        elif line.startswith('# '):
            h = doc.add_paragraph(line.replace('# ', ''), style='Heading 1')
        # Bullets
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        # Standard with Bold support
        else:
            p = doc.add_paragraph()
            # Handle **bold**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    Adds a hyperlink to a paragraph.
    """
    # This gets access to the document.xml.rels file and gets a new relation id
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add color if it is given
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Add underline if it is given
    if underline:
        u = OxmlElement('w:underline')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def create_docx_package(md_text: str, renaming_map: dict):
    """Enhanced docx generator with hyperlinks for renamed files"""
    doc = DocxDocument()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)

    # Prepare search pattern for filenames (Original filenames)
    # We escape them to be safe in regex
    sorted_orig_names = sorted(renaming_map.keys(), key=len, reverse=True)
    
    lines = md_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith('### '):
            p = doc.add_paragraph('', style='Heading 3')
            text = line.replace('### ', '')
        elif line.startswith('## '):
            p = doc.add_paragraph('', style='Heading 2')
            text = line.replace('## ', '')
        elif line.startswith('# '):
            p = doc.add_paragraph('', style='Heading 1')
            text = line.replace('# ', '')
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph('', style='List Bullet')
            text = line[2:]
        else:
            p = doc.add_paragraph()
            text = line

        # Process the text for bolding AND hyperlinks
        # First, handle bolding to split into chunks
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            is_bold = part.startswith('**') and part.endswith('**')
            actual_text = part[2:-2] if is_bold else part
            
            # Now, in each chunk, search for original filenames and replace with linked new names
            ptr = 0
            while ptr < len(actual_text):
                found = False
                for orig in sorted_orig_names:
                    if actual_text.startswith(orig, ptr):
                        new_name = renaming_map[orig]
                        run = p.add_run(actual_text[ptr:ptr]) # empty run to anchor
                        add_hyperlink(p, new_name, new_name)
                        ptr += len(orig)
                        found = True
                        break
                if not found:
                    run = p.add_run(actual_text[ptr:ptr+1])
                    if is_bold: run.bold = True
                    ptr += 1

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

UPLOAD_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "data", "uploads"))
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Dependency to get DB session
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@app.get("/api/schema")
def get_schemas():
    tax_path = os.path.join(os.path.dirname(__file__), "resources", "taxonomy.json")
    field_path = os.path.join(os.path.dirname(__file__), "resources", "field_schema.json")
    
    with open(tax_path, "r", encoding="utf-8") as f:
        taxonomy = json.load(f)
    with open(field_path, "r", encoding="utf-8") as f:
        field_schema = json.load(f)
        
    return {"taxonomy": taxonomy, "field_schema": field_schema}

@app.get("/api/descriptions")
def get_descriptions():
    desc_path = os.path.join(os.path.dirname(__file__), "resources", "descriptions.json")
    with open(desc_path, "r", encoding="utf-8") as f:
        return json.load(f)

@app.post("/api/descriptions")
async def save_descriptions(request: Request):
    desc_path = os.path.join(os.path.dirname(__file__), "resources", "descriptions.json")
    data = await request.json()
    with open(desc_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return {"status": "success"}

from pydantic import BaseModel

class APIKeyPayload(BaseModel):
    api_key: str

@app.get("/api/settings/apikey")
def get_api_key():
    from dotenv import dotenv_values
    env_path = os.path.join(os.path.dirname(__file__), ".env")
    env_config = dotenv_values(env_path) if os.path.exists(env_path) else {}
    key = env_config.get("OPENAI_API_KEY", "")
    masked = key[:6] + "*"*(len(key)-10) + key[-4:] if len(key) > 10 else ""
    return {"masked_key": masked, "has_key": bool(key)}

@app.post("/api/settings/apikey")
def set_api_key(payload: APIKeyPayload):
    env_path = os.path.join(os.path.dirname(__file__), ".env")
    new_key = payload.api_key.strip().replace('"', '').replace("'", "")
    lines = []
    if os.path.exists(env_path):
        with open(env_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
    
    with open(env_path, "w", encoding="utf-8") as f:
        found = False
        for line in lines:
            if line.startswith("OPENAI_API_KEY="):
                f.write(f"OPENAI_API_KEY={new_key}\n")
                found = True
            else:
                f.write(line)
        if not found:
            f.write(f"OPENAI_API_KEY={new_key}\n")
    return {"status": "success"}

from fastapi import BackgroundTasks

def run_extraction_task(doc_id: int, file_path: str, filename: str, schema: str, ai_context: str):
    db = SessionLocal()
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc: return
        
        doc.status = "processing"
        db.commit()
        
        extracted_json = process_pdf_with_vision(
            pdf_path=file_path, 
            original_filename=filename,
            schema_json_str=schema, 
            ai_context=ai_context
        )
        
        meta = extracted_json.get("metadata", {})
        doc.label = meta.get("label", "UNKNOWN")
        doc.category = meta.get("category", "UNKNOWN")
        doc.extracted_data = extracted_json
        doc.status = "completed"
        db.commit()
    except Exception as e:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if doc:
            doc.status = "failed"
            doc.extracted_data = {"error": str(e)}
            db.commit()
    finally:
        db.close()

@app.post("/api/convert")
async def extract_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...), 
    extraction_schema: str = Form("{}"), 
    ai_context: str = Form(""), 
    pratica: str = Form(None),
    db: Session = Depends(get_db)
):
    print(f"[SERVER] Elaborazione: {file.filename}")
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")
        
    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    # Create initial record
    doc_record = Document(
        filename=file.filename,
        file_path=file_path,
        pratica_name=pratica,
        status="pending"
    )
    db.add(doc_record)
    db.commit()
    db.refresh(doc_record)
    
    background_tasks.add_task(
        run_extraction_task, 
        doc_record.id, file_path, file.filename, extraction_schema, ai_context
    )
    
    return {
        "status": "pending",
        "document_id": doc_record.id
    }

@app.get("/api/search")
def search_documents(q: str, db: Session = Depends(get_db)):
    # Simple multi-column search
    results = db.query(Document).filter(
        (Document.filename.like(f"%{q}%")) |
        (Document.pratica_name.like(f"%{q}%")) |
        (Document.label.like(f"%{q}%")) |
        (Document.category.like(f"%{q}%"))
    ).order_by(Document.created_at.desc()).all()
    
    return [
        {
            "id": d.id, 
            "filename": d.filename, 
            "pratica_name": d.pratica_name,
            "label": d.label, 
            "category": d.category,
            "status": d.status,
            "created_at": d.created_at
        } for d in results
    ]

@app.get("/api/documents/{doc_id}")
def get_document_status(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc: raise HTTPException(404)
    return {
        "id": doc.id,
        "status": doc.status,
        "extracted_data": doc.extracted_data,
        "label": doc.label,
        "category": doc.category
    }

@app.post("/api/convert-server-path")
async def extract_document_server_path(
    background_tasks: BackgroundTasks,
    server_filename: str = Form(...),
    schema: str = Form("{}"), 
    ai_context: str = Form(""), 
    pratica: str = Form(None),
    db: Session = Depends(get_db)
):
    file_path = os.path.join(UPLOAD_DIR, server_filename)
    if not os.path.exists(file_path):
        raise HTTPException(404, "Server file not found.")
        
    doc_record = Document(
        filename=server_filename,
        file_path=file_path,
        pratica_name=pratica,
        status="pending"
    )
    db.add(doc_record)
    db.commit()
    db.refresh(doc_record)

    background_tasks.add_task(
        run_extraction_task, 
        doc_record.id, file_path, server_filename, schema, ai_context
    )
    
    return {
        "status": "pending",
        "document_id": doc_record.id
    }

@app.get("/api/pratica/{name}")
def get_pratica(name: str, db: Session = Depends(get_db)):
    docs = db.query(Document).filter(Document.pratica_name == name).all()
    if not docs:
        raise HTTPException(status_code=404, detail="Nessuna pratica trovata.")
        
    payload = []
    for doc in docs:
        payload.append({
            "id": doc.id,
            "filename": doc.filename,
            "category": doc.category,
            "label": doc.label,
            "extracted_data": doc.extracted_data,
            "created_at": doc.created_at
        })
        
    return {"pratica_name": name, "documents": payload}

@app.get("/api/pratiche")
def list_pratiche(db: Session = Depends(get_db)):
    pratiche = db.query(Document.pratica_name).filter(Document.pratica_name.isnot(None)).distinct().all()
    return {"pratiche": [p[0] for p in pratiche if p[0]]}

@app.get("/api/documents")
def list_documents(db: Session = Depends(get_db)):
    docs = db.query(Document).order_by(Document.created_at.desc()).all()
    return [{"id": d.id, "filename": d.filename, "label": d.label, "category": d.category} for d in docs]

from fastapi.responses import FileResponse
@app.get("/api/documents/{doc_id}/pdf")
def get_pdf(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc or not os.path.exists(doc.file_path):
        raise HTTPException(404, "File non trovato nei record archiviati.")
    # Content-Disposition "inline" allows browser to render in iframe/modal.
    # We remove the filename parameter which forces "attachment".
    return FileResponse(doc.file_path, media_type="application/pdf")

@app.delete("/api/pratica/{name}")
def delete_pratica(name: str, db: Session = Depends(get_db)):
    docs = db.query(Document).filter(Document.pratica_name == name).all()
    if not docs:
        raise HTTPException(404, "Pratica non trovata.")
    
    # Optional: Delete physical files
    for doc in docs:
        if os.path.exists(doc.file_path):
            try:
                os.remove(doc.file_path)
            except:
                pass # Silently fail if locked
    
    db.query(Document).filter(Document.pratica_name == name).delete()
    db.commit()
    return {"status": "success", "message": f"Pratica '{name}' eliminata completamente."}

@app.post("/api/pratica/{name}/cross-analyze")
async def cross_analyze_pratica(name: str, rules: str = Form(...), db: Session = Depends(get_db)):
    docs = db.query(Document).filter(Document.pratica_name == name, Document.status == "completed").all()
    if not docs:
        raise HTTPException(404, "Nessun documento analizzato trovato per questa pratica.")
    
    # Extract data from docs
    data_list = []
    for d in docs:
        data_list.append({
            "filename": d.filename,
            "data": d.extracted_data
        })
    
    # Run analysis
    analysis_text = perform_cross_analysis(data_list, rules)
    
    # Save result
    record = CrossAnalysis(
        pratica_name=name,
        rules_prompt=rules,
        analysis_result=analysis_text
    )
    db.add(record)
    db.commit()
    db.refresh(record)
    
    return {
        "id": record.id,
        "pratica_name": name,
        "result": analysis_text,
        "created_at": record.created_at
    }

from fastapi.responses import StreamingResponse

@app.get("/api/analysis/{analysis_id}/export-word")
def export_analysis_word(analysis_id: int, db: Session = Depends(get_db)):
    record = db.query(CrossAnalysis).filter(CrossAnalysis.id == analysis_id).first()
    if not record:
        raise HTTPException(404, "Analisi non trovata.")
    
    filename = f"Analisi_{record.pratica_name}_{record.created_at.strftime('%Y%m%d_%H%M')}.docx"
    
    doc_stream = create_docx_from_markdown(record.analysis_result)
    
    return StreamingResponse(
        doc_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/api/pratica/{name}/cross-analyses")
async def list_cross_analyses(name: str, db: Session = Depends(get_db)):
    rows = db.query(CrossAnalysis).filter(CrossAnalysis.pratica_name == name).order_by(CrossAnalysis.created_at.desc()).all()
    return [
        {
            "id": r.id,
            "rules": r.rules_prompt,
            "result": r.analysis_result,
            "created_at": r.created_at
        } for r in rows
    ]

@app.get("/api/presets")
async def get_presets(db: Session = Depends(get_db)):
    return db.query(AnalysisPreset).all()

@app.post("/api/presets")
async def save_preset(data: dict, db: Session = Depends(get_db)):
    # data: {name: "", rules: [...]}
    name = data.get("name")
    rules = data.get("rules")
    if not name or not rules:
        raise HTTPException(400, "Nome e regole obbligatori.")
    
    # Update if exists
    existing = db.query(AnalysisPreset).filter(AnalysisPreset.name == name).first()
    if existing:
        existing.rules_json = rules
    else:
        new_preset = AnalysisPreset(name=name, rules_json=rules)
        db.add(new_preset)
    
    db.commit()
    return {"status": "ok"}

@app.delete("/api/presets/{preset_id}")
async def delete_preset(preset_id: int, db: Session = Depends(get_db)):
    preset = db.query(AnalysisPreset).filter(AnalysisPreset.id == preset_id).first()
    if preset:
        db.delete(preset)
        db.commit()
    return {"status": "ok"}

@app.post("/api/split")
async def split_document(file: UploadFile = File(...), chunk_size: int = Form(None), ranges: str = Form(None)):
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(400, "Only PDF files supported for splitting.")
        
    temp_path = os.path.join(UPLOAD_DIR, f"temp_split_{file.filename}")
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
        
    try:
        doc = fitz.open(temp_path)
        base_name = file.filename.replace(".pdf", "")
        split_files = []
        num_pages = len(doc)

        split_targets = [] # List of (start, end) 0-indexed
        
        if ranges:
            # Parse "1-2, 3-3, 4-10"
            parts = [p.strip() for p in ranges.split(',')]
            for p in parts:
                if '-' in p:
                    s_str, e_str = p.split('-')
                    s, e = int(s_str.strip()) - 1, int(e_str.strip()) - 1
                    split_targets.append((max(0, s), min(num_pages - 1, e)))
                else:
                    idx = int(p.strip()) - 1
                    split_targets.append((max(0, idx), min(num_pages - 1, idx)))
        elif chunk_size:
            for i in range(0, num_pages, chunk_size):
                end_page = min(i + chunk_size - 1, num_pages - 1)
                split_targets.append((i, end_page))
        else:
            # Default to every page if nothing provided
            for i in range(num_pages):
                split_targets.append((i, i))

        for idx, (s, e) in enumerate(split_targets):
            new_filename = f"{base_name}_taglio_{idx + 1}.pdf"
            new_path = os.path.join(UPLOAD_DIR, new_filename)
            new_doc = fitz.open()
            new_doc.insert_pdf(doc, from_page=s, to_page=e)
            new_doc.save(new_path)
            new_doc.close()
            split_files.append(new_filename)
            
        doc.close()
        os.remove(temp_path)
        return {"status": "success", "files": split_files}
    except Exception as e:
        if os.path.exists(temp_path): os.remove(temp_path)
        raise HTTPException(500, str(e))

@app.post("/api/export-excel")
async def export_excel(payload: Request):
    """
    Expects JSON: { 
      "filename": "name.xlsx", 
      "headers": [ { "text": "...", "type": "meta|field|extra" }, ... ], 
      "rows": [ [cell1, cell2, ...], ... ] 
    }
    """
    data = await payload.json()
    headers_config = data.get("headers", [])
    rows = data.get("rows", [])
    output_filename = data.get("filename", "Estrazione.xlsx")

    # Create DataFrame
    header_texts = [h["text"] for h in headers_config]
    df = pd.DataFrame(rows, columns=header_texts)

    # Color Mapping (Hex codes to match UI)
    colors = {
        "meta": "A0A68D",  # Sage
        "field": "CDBE9F", # Tan
        "extra": "5DADE2"  # Azzurro (Light Blue)
    }

    # Write to Excel with formatting
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Estrazioni')
        
        # Access openpyxl objects for styling
        workbook = writer.book
        worksheet = writer.sheets['Estrazioni']
        
        from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
        
        # Style Headers
        thick_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
        
        for col_idx, h_cfg in enumerate(headers_config, 1):
            cell = worksheet.cell(row=1, column=col_idx)
            h_type = h_cfg.get("type", "meta")
            fill_color = colors.get(h_type, "FFFFFF")
            
            cell.fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid")
            cell.font = Font(color="FFFFFF", bold=True, name="Arial", size=11) if h_type != "extra" else Font(color="FFFFFF", bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thick_border
            
            # Adjust column width
            column_letter = cell.column_letter
            worksheet.column_dimensions[column_letter].width = 25

    output.seek(0)
    
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={output_filename}"}
    )

@app.post("/api/pdf-info")
async def get_pdf_info(file: UploadFile = File(...)):
    temp_path = os.path.join(UPLOAD_DIR, f"info_{file.filename}")
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    try:
        doc = fitz.open(temp_path)
        pages = len(doc)
        doc.close()
        os.remove(temp_path)
        return {"pages": pages}
    except Exception as e:
        if os.path.exists(temp_path): os.remove(temp_path)
        raise HTTPException(500, str(e))

@app.get("/api/pratica/{name}/full-package")
def get_full_package(name: str, db: Session = Depends(get_db)):
    try:
        docs = db.query(Document).filter(Document.pratica_name == name, Document.status == "completed").all()
        if not docs:
            raise HTTPException(404, "Nessun documento trovato per questa pratica.")
        
        # 1. Renaming Map Logic
        renaming_map = {} # original_filename -> new_filename
        used_names = set()
        
        doc_data = [] # List of {doc, new_name}
        
        for d in docs:
            base = os.path.splitext(d.filename)[0]
            ext = os.path.splitext(d.filename)[1]
            cat = d.category or "Generale"
            subcat = d.label or "Nessuna"
            
            # Clean names for filesystem compatibility
            clean_cat = re.sub(r'[\\/*?:"<>|]', "_", cat)
            clean_sub = re.sub(r'[\\/*?:"<>|]', "_", subcat)
            
            new_name_base = f"{base}_{clean_cat}_{clean_sub}"
            
            # Conflict Resolution
            final_name = f"{new_name_base}{ext}"
            counter = 1
            while final_name in used_names:
                final_name = f"{new_name_base}_v{counter}{ext}"
                counter += 1
            
            used_names.add(final_name)
            renaming_map[d.filename] = final_name
            doc_data.append({"doc": d, "new_name": final_name})

        # 2. Latest Cross Analysis
        last_analysis = db.query(CrossAnalysis).filter(CrossAnalysis.pratica_name == name).order_by(CrossAnalysis.created_at.desc()).first()
        
        # 3. ZIP Creation
        temp_zip = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
        temp_zip_path = temp_zip.name
        temp_zip.close()
        
        with zipfile.ZipFile(temp_zip_path, 'w') as zf:
            # A. Add PDFs
            for item in doc_data:
                d = item["doc"]
                if os.path.exists(d.file_path):
                    zf.write(d.file_path, item["new_name"])
            
            # B. Add Excel
            all_field_keys = set()
            for item in doc_data:
                fields = item["doc"].extracted_data.get("fields", {})
                for k in fields.keys(): all_field_keys.add(k)
            field_list = sorted(list(all_field_keys))
            
            headers = ["Nuovo Nome File", "Categoria", "Sottocategoria", "Data Atto", "Data Prot."] + field_list
            excel_rows = []
            for item in doc_data:
                d = item["doc"]
                meta = d.extracted_data.get("metadata", {})
                fields = d.extracted_data.get("fields", {})
                row = [
                    item["new_name"],
                    d.category,
                    d.label,
                    meta.get("data_documento", "-"),
                    meta.get("data_protocollo", "-")
                ]
                for f in field_list:
                    entry = fields.get(f, "-")
                    val = entry.get("value", entry) if isinstance(entry, dict) else entry
                    row.append(val)
                excel_rows.append(row)
                
            df = pd.DataFrame(excel_rows, columns=headers)
            excel_stream = io.BytesIO()
            with pd.ExcelWriter(excel_stream, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Archivio_Dati')
                ws = writer.sheets['Archivio_Dati']
                from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
                
                sage_fill = PatternFill(start_color="A0A68D", end_color="A0A68D", fill_type="solid")
                tan_fill = PatternFill(start_color="CDBE9F", end_color="CDBE9F", fill_type="solid")
                white_font = Font(color="FFFFFF", bold=True)
                
                for i, h in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=i)
                    cell.fill = sage_fill if i <= 5 else tan_fill
                    cell.font = white_font
                    ws.column_dimensions[cell.column_letter].width = 25
                
                for row_idx in range(2, len(excel_rows) + 2):
                    cell = ws.cell(row=row_idx, column=1)
                    filename = cell.value
                    cell.hyperlink = filename
                    cell.font = Font(color="0000FF", underline="single")
                    
            excel_stream.seek(0)
            zf.writestr(f"Estrazione_Dati_{name}.xlsx", excel_stream.getvalue())
            
            if last_analysis:
                word_stream = create_docx_package(last_analysis.analysis_result, renaming_map)
                zf.writestr(f"Audit_Legale_{name}.docx", word_stream.getvalue())
            else:
                fallback_doc = DocxDocument()
                fallback_doc.add_heading(f"Report Archiviazione: {name}", 0)
                fallback_doc.add_paragraph("Nota: Cross-Analisi non ancora eseguita o non trovata.")
                for orig, new in renaming_map.items():
                    p = fallback_doc.add_paragraph(f"- {orig} → ")
                    add_hyperlink(p, new, new)
                fb_stream = io.BytesIO()
                fallback_doc.save(fb_stream)
                zf.writestr(f"Riepilogo_Archiviazione_{name}.docx", fb_stream.getvalue())

        filename = f"Pacchetto_{name}_{datetime.now().strftime('%Y%m%d_%H%M')}.zip"
        
        def iterfile():
            with open(temp_zip_path, mode="rb") as f:
                yield from f
            try:
                os.remove(temp_zip_path)
            except:
                pass

        return StreamingResponse(
            iterfile(),
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"[ERRORE PACCHETTO]\n{traceback.format_exc()}")
        raise HTTPException(500, f"Errore generazione pacchetto: {str(e)}")

# --- SERVE FRONTEND ---
FRONTEND_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "frontend"))

@app.get("/")
async def read_index():
    return FileResponse(os.path.join(FRONTEND_DIR, "index.html"))

# Mount other pages specifically if needed, or just mount the whole dir
app.mount("/", StaticFiles(directory=FRONTEND_DIR), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)
